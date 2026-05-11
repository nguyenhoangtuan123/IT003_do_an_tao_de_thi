from docx import Document

def create_sample():
    doc = Document()
    doc.add_heading('Đề thi mẫu', 0)

    # Câu 1
    doc.add_paragraph('Câu 1: Thủ đô của Việt Nam là gì?')
    doc.add_paragraph('A. Hà Nội')
    doc.add_paragraph('B. Hồ Chí Minh')
    doc.add_paragraph('C. Đà Nẵng')
    doc.add_paragraph('D. Hải Phòng')
    doc.add_paragraph('Đáp án: A')
    doc.add_paragraph('Mức độ: Dễ')
    doc.add_paragraph('Chủ đề: Địa lý')
    doc.add_paragraph('')

    # Câu 2
    doc.add_paragraph('Câu 2: 1 + 1 bằng mấy?')
    doc.add_paragraph('A. 1')
    doc.add_paragraph('B. 2')
    doc.add_paragraph('C. 3')
    doc.add_paragraph('D. 4')
    doc.add_paragraph('Đáp án: B')
    doc.add_paragraph('Mức độ: Dễ')
    doc.add_paragraph('Chủ đề: Toán học')
    doc.add_paragraph('')

    doc.save('sample_upload.docx')
    print("Đã tạo file sample_upload.docx")

if __name__ == "__main__":
    create_sample()
