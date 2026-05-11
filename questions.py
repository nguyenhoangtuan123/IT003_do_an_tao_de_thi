import json
from docx import Document
import copy
class Question:
    def __init__(self, id, question, options, answer,difficulty,topic):
        self.id = id
        self.question = question
        self.options = options
        self.answer = answer
        self.difficulty = difficulty
        self.topic = topic
    def to_dict(self):
        return {
            'id': self.id,
            'question': self.question,
            'options': self.options,
            'answer': self.answer,
            'difficulty': self.difficulty,
            'topic': self.topic
        }
    @staticmethod
    def from_dict(data):
        return Question(
            id=data['id'],
            question=data['question'],
            options=data['options'],
            answer=data['answer'],
            difficulty=data['difficulty'],
            topic=data['topic']
        )
question = [
    Question(1, "What is the capital of France?", 
             {"A": "Paris", "B": "London", "C": "Berlin", "D": "Madrid"}, 
             "A", 
             "Easy", 
             "Geography"),
    Question(2, "What is the largest planet in our solar system?", 
             {"A": "Earth", "B": "Mars", "C": "Jupiter", "D": "Saturn"},
             "C", 
             "Medium", 
             "Science"),
    Question(3, "Who wrote 'To Kill a Mockingbird'?",
             {"A": "William Shakespeare", "B": "Jane Austen", "C": "Harper Lee", "D": "Mark Twain"},
             "C",
             "Hard",
             "Literature"),
    Question(4, "What is the chemical symbol for water?",
             {"A": "H2O", "B": "CO2", "C": "O2", "D": "NaCl"},
             "A",
             "Easy",
             "Science")
]
def save_questions_to_file(questions, filename="questions.json"):
    data = [question.to_dict() for question in questions]
    with open(filename, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=4)
    print(f"Questions saved to {filename}")

def load_questions_from_file(filename="questions.json"):
    with open(filename, 'r', encoding='utf-8') as f:
        data = json.load(f)
    questions = [Question.from_dict(item) for item in data]
    for question in questions:
        print(f"ID: {question.id}")
        print(f"Question: {question.question}")
        for option_key, option_value in question.options.items():
            print(f"{option_key}: {option_value}")
        print(f"Answer: {question.answer}")
        print(f"Difficulty: {question.difficulty}")
        print(f"Topic: {question.topic}")
        print("-" * 40)
    return questions

# Mix questions
import random
def mix_questions(questionset):
    questions = copy.deepcopy(questionset)
    easy = [question for question in questions if question.difficulty.strip().lower() == "easy"]
    medium = [question for question in questions if question.difficulty.strip().lower() == "medium"]
    hard = [question for question in questions if question.difficulty.strip().lower() == "hard"]

    num_easy = len(easy)
    num_medium = len(medium)    
    num_hard = len(hard)
    
    selected = []
    selected += random.sample(easy, num_easy)
    selected += random.sample(medium, num_medium)
    selected += random.sample(hard, num_hard)
    
    random.shuffle(selected)
    for question in selected:
        ket_qua = question.options[question.answer]
        answers = [q for q in question.options.values()]
        random.shuffle(answers)
        new_options = {}
        for i, key in enumerate(question.options.keys()):
            new_options[key] = answers[i]
        question.options = new_options
        question.answer = [key for key, value in question.options.items() if value == ket_qua][0]
    return selected 

def export_questions_to_docx(questions,exam_code, filename="questions.docx"):    
    doc = Document()
    doc.add_heading('De thi', level=0)
    doc.add_paragraph(f"Ma de: {exam_code}")
    for i, question in enumerate(questions):
        doc.add_paragraph(f"Cau {i+1}: {question.question}")
        doc.add_paragraph(f"A: {question.options['A']}")
        doc.add_paragraph(f"B: {question.options['B']}")
        doc.add_paragraph(f"C: {question.options['C']}")
        doc.add_paragraph(f"D: {question.options['D']}")
        doc.add_paragraph(" " * 20)
    doc.save(filename)

if __name__ == "__main__":
    from database import load_questions_from_mongodb, save_questions_to_mongodb
    
    save_questions_to_file(question)
    
    # Thử nghiệm lưu câu hỏi mẫu vào MongoDB
    try:
        save_questions_to_mongodb(question)
        loaded_from_db = load_questions_from_mongodb()
        print(f"Đã tải {len(loaded_from_db)} câu hỏi từ MongoDB.")
    except Exception as e:
        print(f"Lỗi khi tương tác MongoDB: {e}")

    mix = mix_questions(question)
    export_questions_to_docx(mix, "123456")